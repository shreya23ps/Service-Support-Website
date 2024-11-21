# main/views.py
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, JsonResponse
from reportlab.pdfgen import canvas
from io import BytesIO
from docx import Document
from django.contrib import messages

from django.shortcuts import render, get_object_or_404, redirect
from .models import BlogPost, Comment, Contact, Person, PanCardApplication,ElectoralRoll, DeathCertificateApplication, AddressChangeApplication, DrivingLicenseApplication, RationCardApplication, VoterRegistration

def index(request):
    if request.method == 'POST':
        # Handle form submission
        name = request.POST.get('name')
        email = request.POST.get('email')
        subject = request.POST.get('subject')
        message = request.POST.get('message')

        # Create a new contact entry in the database
        contact = Contact(name=name, email=email, subject=subject, message=message)
        contact.save()
        
        return JsonResponse({'status': 'success', 'message': 'Your message has been sent. Thank you!'})
    
    # Render the index.html for GET requests
    return render(request, 'index.html')

def about(request):
    return render(request, 'about.html')

def services(request):
    return render(request, 'services.html')

def portfolio(request):
    return render(request, 'portfolio-details.html')

def team(request):
    return render(request, 'team.html')

def blog(request):
    blogs = BlogPost.objects.all()  # Fetch all blog posts
    print(f"Number of blogs fetched: {blogs.count()}")
    context = {
        'blogs': blogs,
    }
    
    return render(request, 'blog.html', context)  # Update the template path as necessary

def blog_details(request, blog_id):

    # Ensure blog_id is an integer
    blog_post = get_object_or_404(BlogPost, id=blog_id)
    comments = blog_post.comments.all()

    if request.method == "POST":
        name = request.POST.get('name')
        email = request.POST.get('email')
        content = request.POST.get('content')

        # Create and save a new comment
        new_comment = Comment(blog_post=blog_post, name=name, email=email, content=content)
        new_comment.save()

        # Re-render the template with updated comments
        return render(request, 'blog-details.html', {'blog_post': blog_post, 'comments': comments})

    return render(request, 'blog-details.html', {'blog_post': blog_post, 'comments': comments})

def add_blog(request):
    if request.method == 'POST':
        # Get form data from POST request
        title = request.POST['title']
        author = request.POST['author']
        content = request.POST['content']
        
        # Create a new BlogPost object and save it to the database
        new_blog = BlogPost(title=title, author=author, content=content)
        new_blog.save()
        
        # After saving, redirect to the blog listing or another page
        return redirect('blog')  # Replace 'blog_list' with the name of the page you want to redirect to

    # If the request is GET, just render the form
    return render(request, 'add_blog.html')

   
def service_details(request):
    return render(request, 'service-details.html')

def death_application_submitted(request):
     if request.method == 'POST':
        # Retrieving form data
        deceased_name = request.POST.get('deceased_name')
        deceased_dob = request.POST.get('deceased_dob')
        deceased_dod = request.POST.get('deceased_dod')
        deceased_gender = request.POST.get('deceased_gender')
        deceased_aadhaar = request.POST.get('deceased_aadhaar')
        deceased_address = request.POST.get('deceased_address')

        applicant_name = request.POST.get('applicant_name')
        applicant_relation = request.POST.get('applicant_relation')
        applicant_address = request.POST.get('applicant_address')
        applicant_contact = request.POST.get('applicant_contact')
        applicant_email = request.POST.get('applicant_email')

        documents = request.FILES.get('documents', None)

        # Create the death certificate application entry
        death_application = DeathCertificateApplication.objects.create(
            deceased_name=deceased_name,
            deceased_dob=deceased_dob,
            deceased_dod=deceased_dod,
            deceased_gender=deceased_gender,
            deceased_aadhaar=deceased_aadhaar,
            deceased_address=deceased_address,
            applicant_name=applicant_name,
            applicant_relation=applicant_relation,
            applicant_address=applicant_address,
            applicant_contact=applicant_contact,
            applicant_email=applicant_email,
            documents=documents
        )

        # Process the person record
        try:
            person = Person.objects.get(aadhaar_number=deceased_aadhaar)
            if person.is_alive:
                # Mark as deceased
                person.is_alive = False
                person.save()

                # Remove from electoral rolls
                ElectoralRoll.objects.filter(person=person).delete()

                message = 'Death certificate processed successfully. Person marked as deceased and removed from electoral rolls.'
            else:
                message = 'Person is already marked as deceased.'
        except Person.DoesNotExist:
            message = 'Aadhaar number not found in the centralized database.'

        # Show success or error message
        messages.success(request, message)
        return render(request,'death_application_submitted.html')

def death_application(request):
    return render(request,'death_application.html')

def pan_application(request):
    return render(request,'pan_application.html')

def pan_application_submitted(request):
     if request.method == 'POST':
        # Retrieving form data
        applicant_name = request.POST.get('applicant_name')
        applicant_father_name = request.POST.get('applicant_father_name')
        applicant_dob = request.POST.get('applicant_dob')
        applicant_gender = request.POST.get('applicant_gender')
        applicant_aadhaar = request.POST.get('applicant_aadhaar')
        applicant_address = request.POST.get('applicant_address')
        applicant_email = request.POST.get('applicant_email')
        applicant_contact = request.POST.get('applicant_contact')
        id_proof = request.FILES.get('id_proof')
        address_proof = request.FILES.get('address_proof')
        documents = request.FILES.get('documents', None)  # Optional field

        # Check if person exists in Person table
        person, created = Person.objects.get_or_create(
            aadhaar_number=applicant_aadhaar,
            defaults={
                'name': applicant_name,
                'surname': applicant_father_name,
                'is_alive': True  # Assuming person is alive when applying for a PAN card
            }
        )

        if created:
            # If a new person, create an electoral roll entry
            from .models import ElectoralRoll
            ElectoralRoll.objects.create(person=person)

        # Create the PAN card application entry in PanCardApplication
        pan_application = PanCardApplication.objects.create(
            applicant_name=applicant_name,
            applicant_father_name=applicant_father_name,
            applicant_dob=applicant_dob,
            applicant_gender=applicant_gender,
            applicant_aadhaar=applicant_aadhaar,
            applicant_address=applicant_address,
            applicant_email=applicant_email,
            applicant_contact=applicant_contact,
            id_proof=id_proof,
            address_proof=address_proof,
            documents=documents
        )

        # Success message to display
        messages.success(request, 'PAN Card application submitted successfully!')

        return render(request,'pan_application_submitted.html')

def driving_license_application(request):
    return render(request,'driving_license_application.html')

def driving_license_application_submitted(request):
    if request.method == 'POST':
        # Retrieving form data
        applicant_name = request.POST.get('applicant_name')
        applicant_father_name = request.POST.get('applicant_father_name')
        applicant_dob = request.POST.get('applicant_dob')
        applicant_gender = request.POST.get('applicant_gender')
        applicant_aadhaar = request.POST.get('applicant_aadhaar')
        applicant_address = request.POST.get('applicant_address')

        # Retrieving license details
        license_type = request.POST.get('license_type')
        vehicle_class = request.POST.get('vehicle_class')

        # Retrieving contact details
        applicant_contact = request.POST.get('applicant_contact')
        applicant_email = request.POST.get('applicant_email')

        # Retrieving supporting documents
        documents = request.FILES.get('documents')

        # Create the Driving License Application entry
        driving_license_application = DrivingLicenseApplication.objects.create(
            applicant_name=applicant_name,
            applicant_father_name=applicant_father_name,
            applicant_dob=applicant_dob,
            applicant_gender=applicant_gender,
            applicant_aadhaar=applicant_aadhaar,
            applicant_address=applicant_address,
            license_type=license_type,
            vehicle_class=vehicle_class,
            applicant_contact=applicant_contact,
            applicant_email=applicant_email,
            documents=documents
        )

        # Success message to display
        messages.success(request, 'Driving License application submitted successfully!')
    return render(request,'driving_license_application_submitted.html')

def ration_card_application(request):
    return render(request,'ration_card_application.html')

def ration_card_application_submitted(request):
    if request.method == 'POST':
        # Extract data from the POST request
        applicant_name = request.POST.get('applicant_name')
        applicant_father_name = request.POST.get('applicant_father_name')
        applicant_dob = request.POST.get('applicant_dob')
        applicant_gender = request.POST.get('applicant_gender')
        applicant_aadhaar = request.POST.get('applicant_aadhaar')
        applicant_address = request.POST.get('applicant_address')
        applicant_state = request.POST.get('applicant_state')
        applicant_district = request.POST.get('applicant_district')
        applicant_city = request.POST.get('applicant_city')
        applicant_pincode = request.POST.get('applicant_pincode')
        applicant_email = request.POST.get('applicant_email')
        applicant_contact = request.POST.get('applicant_contact')
        family_members = request.POST.get('family_members')
        ration_card_type = request.POST.get('ration_card_type')

        # Handle file uploads (e.g., id_proof, address_proof, and photo)
        id_proof = request.FILES.get('id_proof')
        address_proof = request.FILES.get('address_proof')
        photo = request.FILES.get('photo')

        # Save data to the model
        ration_card_application = RationCardApplication(
            applicant_name=applicant_name,
            applicant_father_name=applicant_father_name,
            applicant_dob=applicant_dob,
            applicant_gender=applicant_gender,
            applicant_aadhaar=applicant_aadhaar,
            applicant_address=applicant_address,
            applicant_state=applicant_state,
            applicant_district=applicant_district,
            applicant_city=applicant_city,
            applicant_pincode=applicant_pincode,
            applicant_email=applicant_email,
            applicant_contact=applicant_contact,
            family_members=family_members,
            ration_card_type=ration_card_type,
            id_proof=id_proof,
            address_proof=address_proof,
            photo=photo,
        )
        ration_card_application.save()
        
    return render(request,'ration_card_application_submitted.html')

def address_change_application(request):
    return render(request,'address_change_application.html')

def address_change_application_submitted(request):
    if request.method == 'POST':
        # Retrieving form data
        applicant_name = request.POST.get('applicant_name')
        applicant_dob = request.POST.get('applicant_dob')
        applicant_gender = request.POST.get('applicant_gender')
        applicant_aadhaar = request.POST.get('applicant_aadhaar')
        old_address = request.POST.get('old_address')
        new_address = request.POST.get('new_address')

        # Retrieving files
        address_proof = request.FILES.get('address_proof')
        additional_documents = request.FILES.get('additional_documents', None)  # Optional field

        # Create the Address Change Application entry
        address_change_application = AddressChangeApplication.objects.create(
            applicant_name=applicant_name,
            applicant_dob=applicant_dob,
            applicant_gender=applicant_gender,
            applicant_aadhaar=applicant_aadhaar,
            old_address=old_address,
            new_address=new_address,
            address_proof=address_proof,
            additional_documents=additional_documents
        )

        # Success message to display
        messages.success(request, 'Address Change application submitted successfully!')
    return render(request,'address_change_application_submitted.html')

def voter_registration(request):
    return render(request,'voter_registration.html')

def voter_registration_submitted(request):
     if request.method == 'POST':
        # Get data from the POST request
        applicant_name = request.POST.get('applicant_name')
        applicant_dob = request.POST.get('applicant_dob')
        applicant_gender = request.POST.get('applicant_gender')
        applicant_aadhaar = request.POST.get('applicant_aadhaar')
        applicant_address = request.POST.get('applicant_address')
        applicant_pincode = request.POST.get('applicant_pincode')
        applicant_email = request.POST.get('applicant_email')
        applicant_contact = request.POST.get('applicant_contact')
        
        # Handle file uploads
        id_proof = request.FILES.get('id_proof')
        address_proof = request.FILES.get('address_proof')

        # Create and save the VoterRegistration object
        voter_registration = VoterRegistration(
            applicant_name=applicant_name,
            applicant_dob=applicant_dob,
            applicant_gender=applicant_gender,
            applicant_aadhaar=applicant_aadhaar,
            applicant_address=applicant_address,
            applicant_pincode=applicant_pincode,
            applicant_email=applicant_email,
            applicant_contact=applicant_contact,
            id_proof=id_proof,
            address_proof=address_proof,
        )
        voter_registration.save()
        return render(request,'voter_registration_submitted.html')

def check_pan_card_status(request):
    return render(request,'check_pan_card_status.html')

def pan_card_status(request):
    context = {}
    if request.method == 'POST':
        aadhaar_number = request.POST.get('aadhaar_number', '').strip()
        
        if aadhaar_number:
            # Search for the applicant in the Person model using the Aadhaar number
            try:
                person = Person.objects.get(aadhaar_number=aadhaar_number)  # Look up Person by Aadhaar number
                context['person'] = person
            except Person.DoesNotExist:
                context['error'] = 'No record found for this Aadhaar number.'
        else:
            context['error'] = 'Please enter a valid Aadhaar number.'

    return render(request, 'pan_card_status.html', context)


def check_driving_license_status(request):
    return render(request,'check_driving_license_status.html')

def driving_license_status(request):
    context = {}
    if request.method == 'POST':
        aadhaar_number = request.POST.get('aadhaar_number', '').strip()
        
        # Check if Aadhaar number is provided
        if aadhaar_number:
            try:
                # Query the Person model using the Aadhaar number
                person_record = Person.objects.get(aadhaar_number=aadhaar_number)
                
                # Pass only the required fields to the template
                context['name'] = person_record.name
                context['aadhaar_number'] = person_record.aadhaar_number
                context['status'] = 'Active' if person_record.is_alive == 1 else 'Inactive'
            except Person.DoesNotExist:
                context['error'] = 'No record found for this Aadhaar number.'
        else:
            context['error'] = 'Please provide a valid Aadhaar number.'

    return render(request, 'driving_license_status.html', context)

def check_ration_card_status(request):
    return render(request,'check_ration_card_status.html')

def ration_card_status(request):
    context = {}
    if request.method == 'POST':
        aadhaar_number = request.POST.get('aadhaar_number', '').strip()
        
        if aadhaar_number:
            try:
                # Query the Person model using the Aadhaar number
                person_record = Person.objects.get(aadhaar_number=aadhaar_number)
                
                # Pass only the required fields to the template
                context['name'] = person_record.name
                context['aadhaar_number'] = person_record.aadhaar_number
                context['status'] = 'Active' if person_record.is_alive == 1 else 'Inactive'
            except Person.DoesNotExist:
                context['error'] = 'No record found for this Aadhaar number.'
        else:
            context['error'] = 'Please provide a valid Aadhaar number.'

    return render(request, 'ration_card_status.html', context)
    return render(request,'ration_card_status.html')

def check_pension_status(request):
    return render(request,'check_pension_status.html')

def pension_status(request):
    context = {}
    if request.method == 'POST':
        aadhaar_number = request.POST.get('aadhaar_number', '').strip()
        
        if aadhaar_number:
            try:
                # Query the Person model using the Aadhaar number
                person_record = Person.objects.get(aadhaar_number=aadhaar_number)
                
                # Pass only the required fields to the template
                context['name'] = person_record.name
                context['aadhaar_number'] = person_record.aadhaar_number
                context['status'] = 'Eligible' if person_record.is_alive == 1 else 'Not Eligible'
            except Person.DoesNotExist:
                context['error'] = 'No record found for this Aadhaar number.'
        else:
            context['error'] = 'Please provide a valid Aadhaar number.'

    return render(request, 'pension_status.html', context)
    

#code for downloading service list as pdf and doc

# Service list
SERVICES = [
    "Death Certificate Application",
    "PAN Card Application",
    "Driving License Application",
    "Ration Card Application",
    "Address Change Services",
    "Voter Registration Services",
    "Check PAN card status",
    "Check Driving License status",
    "Check Ration Card status",
    "Check Pension status",
]

# View to download PDF
def download_services_pdf(request):
    # Create a file-like buffer to receive PDF data
    buffer = BytesIO()
    
    # Create the PDF object
    p = canvas.Canvas(buffer)
    
    # Draw the services list on the PDF
    p.drawString(100, 800, "Services List")
    y = 750
    for service in SERVICES:
        p.drawString(100, y, service)
        y -= 30
    
    # Close the PDF object cleanly
    p.showPage()
    p.save()
    
    # Get the value of the BytesIO buffer and write it to the response
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="services_list.pdf"'
    
    return response

# View to download DOC
def download_services_doc(request):
    # Create a Word document
    doc = Document()
    doc.add_heading('Services List', 0)
    
    for service in SERVICES:
        doc.add_paragraph(service)
    
    # Create a response with the DOC file
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="services_list.docx"'
    
    # Save the document to the response
    doc.save(response)
    
    return response


# Add additional view functions for other templates

